from docx import Document

def convert_docx_file(tilte_docx):
    doc=Document(tilte_docx)
    for i in doc.paragraphs:
        for run in i.runs:
            if run.bold:
                if 'а' in run.text:
                    run.text=run.text.replace('а','a')
                if 'А' in run.text:
                    run.text=run.text.replace('А','A')
                if 'Н' in run.text:
                    run.text= run.text.replace('Н', 'H')
                if 'о' in run.text:
                    run.text = run.text.replace('о', 'o')
                if 'О' in run.text:
                    run.text = run.text.replace('О', 'O')
                if 'е' in run.text:
                    run.text= run.text.replace('е', 'e')
                if 'Е' in run.text:
                    run.text = run.text.replace('Е', 'E')
                if 'с' in run.text:
                    run.text = run.text.replace('с', 'c')
                if 'С' in run.text:
                    run.text = run.text.replace('С', 'C')
                if 'Т' in run.text:
                    run.text = run.text.replace('Т', 'T')
                if 'і' in run.text:
                    run.text = run.text.replace('і', 'i')
                if 'у' in run.text:
                    run.text = run.text.replace('у', 'y')
                if 'х' in run.text:
                    run.text = run.text.replace('х', 'x')
                if 'р' in run.text:
                    run.text = run.text.replace('р', 'p')
                # if 'к' in run.text:
                #     run.text = run.text.replace('к', 'k')
                if 'К' in run.text:
                    run.text = run.text.replace('К', 'K')
                if 'М' in run.text:
                    run.text = run.text.replace('М', 'M')
                if 'В' in run.text:
                    run.text = run.text.replace('В', 'B')


            else:
                if 'а' in run.text:
                    run.text=run.text.replace('а','a')
                if 'А' in run.text:
                    run.text=run.text.replace('А','A')
                if 'Н' in run.text:
                    run.text= run.text.replace('Н', 'H')
                if 'о' in run.text:
                    run.text = run.text.replace('о', 'o')
                if 'О' in run.text:
                    run.text = run.text.replace('О', 'O')
                if 'е' in run.text:
                    run.text= run.text.replace('е', 'e')
                if 'Е' in run.text:
                    run.text = run.text.replace('Е', 'E')
                if 'с' in run.text:
                    run.text = run.text.replace('с', 'c')
                if 'С' in run.text:
                    run.text = run.text.replace('С', 'C')
                if 'Т' in run.text:
                    run.text = run.text.replace('Т', 'T')
                if 'і' in run.text:
                    run.text = run.text.replace('і', 'i')
                if 'у' in run.text:
                    run.text = run.text.replace('у', 'y')
                if 'х' in run.text:
                    run.text = run.text.replace('х', 'x')
                if 'р' in run.text:
                    run.text = run.text.replace('р', 'p')
                # if 'к' in run.text:
                #     run.text = run.text.replace('к', 'k')
                if 'К' in run.text:
                    run.text = run.text.replace('К', 'K')
                if 'М' in run.text:
                    run.text = run.text.replace('М', 'M')
                if 'В' in run.text:
                    run.text = run.text.replace('В', 'B')
    for table in doc.tables:
        for coll in table.columns:
            for cell in coll.cells:
                for pr in cell.paragraphs:
                    if 'а' in pr.text:
                        pr.text=pr.text.replace('а','a')
                    if 'А' in pr.text:
                        pr.text = pr.text.replace('А', 'A')
                    if 'Н' in pr.text:
                        pr.text = pr.text.replace('Н', 'H')
                    if 'о' in pr.text:
                        pr.text = pr.text.replace('о', 'o')
                    if 'О' in pr.text:
                        pr.text = pr.text.replace('О', 'O')
                    if 'е' in pr.text:
                        pr.text = pr.text.replace('е', 'e')
                    if 'Е' in pr.text:
                        pr.text = pr.text.replace('Е', 'E')
                    if 'с' in pr.text:
                        pr.text = pr.text.replace('с', 'c')
                    if 'С' in pr.text:
                        pr.text = pr.text.replace('С', 'C')
                    if 'Т' in pr.text:
                        pr.text = pr.text.replace('Т', 'T')
                    if 'і' in pr.text:
                        pr.text = pr.text.replace('і', 'i')
                    if 'у' in pr.text:
                        pr.text = pr.text.replace('у', 'y')
                    if 'х' in pr.text:
                        pr.text = pr.text.replace('х', 'x')
                    if 'р' in pr.text:
                        pr.text = pr.text.replace('р', 'p')
                    # if 'к' in pr.text:
                    #     pr.text = pr.text.replace('к', 'k')
                    if 'К' in pr.text:
                        pr.text = pr.text.replace('К', 'K')
                    if 'М' in pr.text:
                        pr.text = pr.text.replace('М', 'M')
                    if 'В' in pr.text:
                        pr.text = pr.text.replace('В', 'B')

    return doc.save(tilte_docx)
    print("Hello world !")
title='Hello.docx'
convert_docx_file(title)
